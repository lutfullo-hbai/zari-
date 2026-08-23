import logging
import re
from datetime import datetime
from pathlib import Path

from skills.base import BaseSkill

log = logging.getLogger("zari")

DOC_EXT_RE = re.compile(r"([\w ./~\-]+\.(pdf|docx?|xlsx?|csv|txt|md))\b", re.IGNORECASE)
MAX_TEXT_CHARS = 20000


class DocumentSkill(BaseSkill):
    """PDF/Word/Excel/CSV/TXT hujjatlarni o'qish va yangi docx/xlsx yaratish."""

    priority = 70
    timeout = 30.0

    async def execute(self, query: str) -> dict | None:
        text = query.lower()
        wants_create = any(w in text for w in ["yarat", "yarat", "yasа", "create"])
        if wants_create:
            return await self._create(query)

        path = self._extract_doc_path(query)
        if not path:
            return None
        if not path.is_file():
            return {
                "response": f"Hujjat topilmadi: {path}",
                "context": "",
                "source": "documents",
            }

        content = self._extract_text(path)
        if content is None:
            return {
                "response": f"Bu formatni o'qa olmayman: {path.suffix}",
                "context": "",
                "source": "documents",
            }
        if content == "":
            return {
                "response": f"{path.name} bo'sh hujjat.",
                "context": "",
                "source": "documents",
            }

        preview = content[:1500]
        response = f"{path.name} ({len(content)} belgi):\n{preview}"
        if len(content) > 1500:
            response += "\n... (qolganini xulosa qilib berishim mumkin)"
        # context to'liq matn — Brain buni ko'rib xulosa qiladi
        return {
            "response": response,
            "context": content[:MAX_TEXT_CHARS],
            "source": "documents",
        }

    def _extract_doc_path(self, query: str) -> Path | None:
        m = DOC_EXT_RE.search(query)
        if m:
            p = Path(m.group(1).strip()).expanduser()
            if not p.is_absolute():
                cand = Path.home() / p
                p = cand if cand.exists() else p
            return p
        return None

    def _extract_text(self, path: Path) -> str | None:
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                from pypdf import PdfReader

                reader = PdfReader(str(path))
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            if ext == ".docx":
                import docx as docx_mod

                d = docx_mod.Document(str(path))
                paras = [p.text for p in d.paragraphs]
                for table in d.tables:
                    for row in table.rows:
                        paras.append("\t".join(c.text for c in row.cells))
                return "\n".join(paras)
            if ext == ".xlsx":
                import openpyxl

                wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
                lines = []
                for ws in wb.worksheets:
                    lines.append(f"[{ws.title}]")
                    for row in ws.iter_rows(values_only=True):
                        vals = ["" if v is None else str(v) for v in row]
                        if any(vals):
                            lines.append("\t".join(vals))
                wb.close()
                return "\n".join(lines)
            if ext == ".xls":
                return None
            return path.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            log.warning("Hujjat o'qilmadi %s: %s", path.name, e)
            return ""

    async def _create(self, query: str) -> dict | None:
        m = re.search(r'"([^"]+)"|\'([^\']+)\'', query)
        title = next((g for g in (m.groups() if m else ()) if g), None)
        now = datetime.now().strftime("%Y%m%d_%H%M%S")

        if any(w in query.lower() for w in ["excel", "xlsx", "jadval"]):
            return await self._create_xlsx(title or f"hujjat_{now}")
        if any(w in query.lower() for w in ["word", "docx", "hujjat", "matn"]):
            return await self._create_docx(title or f"hujjat_{now}", query)
        return None

    async def _create_docx(self, name: str, query: str) -> dict:
        import docx as docx_mod

        out_dir = Path.home() / "Documents"
        out_dir.mkdir(exist_ok=True)
        body = ""
        bm = re.search(r"(?:matn|content)[:\s]+(.+)$", query, re.IGNORECASE | re.DOTALL)
        if bm:
            body = bm.group(1).strip()
        safe_name = re.sub(r"[^\w\- ]", "", name).strip() or "hujjat"
        out = out_dir / f"{safe_name}.docx"

        d = docx_mod.Document()
        d.add_heading(safe_name, level=1)
        if body:
            d.add_paragraph(body)
        d.save(str(out))
        return {
            "response": f"Word hujjat yaratildi: {out}",
            "context": str(out),
            "source": "documents",
        }

    async def _create_xlsx(self, name: str) -> dict:
        import openpyxl

        out_dir = Path.home() / "Documents"
        out_dir.mkdir(exist_ok=True)
        safe_name = re.sub(r"[^\w\- ]", "", name).strip() or "jadval"
        out = out_dir / f"{safe_name}.xlsx"

        wb = openpyxl.Workbook()
        wb.active.title = "Sheet1"
        wb.save(str(out))
        return {
            "response": f"Excel jadval yaratildi: {out}",
            "context": str(out),
            "source": "documents",
        }
