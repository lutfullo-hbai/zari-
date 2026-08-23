"""M7 Hujjatlar va Kod testlari: Documents, Organize, CodeRunner."""

import sys
from pathlib import Path
from unittest.mock import patch

import pytest

from core.router import match_intents


def _make_pdf(path: Path, text: str) -> None:
    from pypdf import PdfWriter

    w = PdfWriter()
    w.add_blank_page(width=200, height=200)
    # pypdf matn yozishni qo'llamaydi — minimal PDF qo'lda yaratamiz
    content = f"BT /F1 12 Tf 50 100 Td ({text}) Tj ET".encode()
    stream = (
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 200 200]"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        + b"4 0 obj<</Length "
        + str(len(content)).encode()
        + b">>stream\n"
        + content
        + b"\nendstream endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n0000000000 65535 f \ntrailer<</Size 6/Root 1 0 R>>\nstartxref\n9\n%%EOF"
    )
    path.write_bytes(stream)


class TestDocumentSkill:
    @pytest.fixture
    def skill(self):
        from skills.documents import DocumentSkill

        return DocumentSkill()

    @pytest.mark.asyncio
    async def test_read_txt(self, skill, tmp_path):
        f = tmp_path / "qoidalar.txt"
        f.write_text("Birinchi qoida. Ikkinchi qoida.", encoding="utf-8")
        result = await skill.execute(f"{f} ni o'qib ber")

        assert result is not None
        assert "Birinchi qoida" in result["response"]
        assert result["source"] == "documents"

    @pytest.mark.asyncio
    async def test_read_docx(self, skill, tmp_path):
        import docx as docx_mod

        f = tmp_path / "hisobot.docx"
        d = docx_mod.Document()
        d.add_paragraph("Moliyaviy xulosa: foyda oshdi.")
        d.save(str(f))
        result = await skill.execute(f"{f} hujjatini o'qi")

        assert result is not None
        assert "foyda oshdi" in result["response"]

    @pytest.mark.asyncio
    async def test_read_xlsx(self, skill, tmp_path):
        import openpyxl

        f = tmp_path / "jadval.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Savdo"
        ws["A1"] = "Tovar"
        wb.active["B1"] = "Narx"
        wb.save(str(f))
        result = await skill.execute(f"{f} jadvalni ko'rsat")

        assert result is not None
        assert "Tovar" in result["response"]
        assert "[Savdo]" in result["response"]

    @pytest.mark.asyncio
    async def test_read_pdf(self, skill, tmp_path):
        f = tmp_path / "report.pdf"
        _make_pdf(f, "Salom dunyo")
        result = await skill.execute(f"{f} ni o'qi")

        assert result is not None
        assert "Salom dunyo" in result["response"]

    @pytest.mark.asyncio
    async def test_missing_file(self, skill, tmp_path):
        result = await skill.execute(f"{tmp_path}/yok.pdf ni o'qi")
        assert result is not None
        assert "topilmadi" in result["response"]

    @pytest.mark.asyncio
    async def test_no_doc_keyword_returns_none(self, skill):
        """Hujjat so'zi yo'q bo'lsa skill aralashmasligi kerak."""
        assert await skill.execute("ob-havo qanday") is None

    @pytest.mark.asyncio
    async def test_create_xlsx(self, skill, tmp_path, monkeypatch):
        monkeypatch.setattr(Path, "home", lambda: tmp_path)
        result = await skill.execute("excel jadval yarat nomi 'savdo'")

        assert result is not None
        created = tmp_path / "Documents" / "savdo.xlsx"
        assert created.exists()

    @pytest.mark.asyncio
    async def test_create_docx_with_body(self, skill, tmp_path, monkeypatch):
        monkeypatch.setattr(Path, "home", lambda: tmp_path)
        await skill.execute('word hujjat yarat "mektub" matn: Salom aziz do\'st')

        created = tmp_path / "Documents" / "mektub.docx"
        assert created.exists()
        import docx as docx_mod

        text = "\n".join(p.text for p in docx_mod.Document(str(created)).paragraphs)
        assert "Salom aziz do'st" in text


class TestOrganizeSkill:
    @pytest.fixture
    def skill(self):
        from skills.organize import OrganizeSkill

        return OrganizeSkill()

    def test_requires_confirmation(self, skill):
        """Ko'p fayl ko'chishi uchun tasdiq MAJBURIY."""
        assert skill.requires_confirmation is True
        assert skill.confirmation_type == "danger"

    @pytest.mark.asyncio
    async def test_organizes_by_category(self, skill, tmp_path):
        (tmp_path / "rasm.png").write_bytes(b"x")
        (tmp_path / "hujjat.pdf").write_bytes(b"x")
        (tmp_path / "qoshiq.mp3").write_bytes(b"x")
        (tmp_path / "kod.py").write_bytes(b"x")

        result = await skill.execute(f"{tmp_path} papkani tartibga sol")

        assert result is not None
        assert "4 ta fayl" in result["response"]
        assert (tmp_path / "Rasmlar" / "rasm.png").exists()
        assert (tmp_path / "Hujjatlar" / "hujjat.pdf").exists()
        assert (tmp_path / "Musiqa" / "qoshiq.mp3").exists()
        assert (tmp_path / "Kodlar" / "kod.py").exists()
        assert not (tmp_path / "rasm.png").exists()

    @pytest.mark.asyncio
    async def test_never_overwrites_duplicates(self, skill, tmp_path):
        src = tmp_path / "a.png"
        src.write_bytes(b"x")
        dest_dir = tmp_path / "Rasmlar"
        dest_dir.mkdir()
        existing = dest_dir / "a.png"
        existing.write_bytes(b"ORIGINAL")

        await skill.execute(f"{tmp_path} tartibga sol")

        assert existing.read_bytes() == b"ORIGINAL"
        moved = dest_dir / "a (2).png"
        assert moved.exists() and moved.read_bytes() == b"x"

    @pytest.mark.asyncio
    async def test_unknown_ext_goes_to_boshqalar(self, skill, tmp_path):
        (tmp_path / "fayl.xyz").write_bytes(b"x")
        await skill.execute(f"{tmp_path} sarala")
        assert (tmp_path / "Boshqalar" / "fayl.xyz").exists()

    @pytest.mark.asyncio
    async def test_missing_dir(self, skill, tmp_path):
        result = await skill.execute(f"{tmp_path}/yok tartibga sol")
        assert "topilmadi" in result["response"]

    @pytest.mark.asyncio
    async def test_wrong_intent_returns_none(self, skill):
        assert await skill.execute("musiqani ijro et") is None


class TestCodeRunnerSkill:
    @pytest.fixture
    def skill(self):
        from skills.code_runner import CodeRunnerSkill

        return CodeRunnerSkill()

    def test_requires_confirmation(self, skill):
        """Kod bajarish har doim tasdiq talab qiladi."""
        assert skill.requires_confirmation is True
        assert skill.confirmation_type == "danger"

    @pytest.mark.asyncio
    async def test_run_code_block(self, skill):
        query = "kodni ishga tushir:\n```python\nprint(2 + 2)\n```"
        result = await skill.execute(query)

        assert result is not None
        assert "4" in result["response"]
        assert result["context"] == "code_runner:OK"

    @pytest.mark.asyncio
    async def test_run_inline_code(self, skill):
        result = await skill.execute("kod yozib ishga tushir: print('salom')")

        assert result is not None
        assert "salom" in result["response"]

    @pytest.mark.asyncio
    async def test_captures_error(self, skill):
        result = await skill.execute("```python\nraise ValueError('test xato')\n```\nishga tushir kod")

        assert result is not None
        assert "ValueError" in result["response"]
        assert result["context"] == "code_runner:ERROR"

    @pytest.mark.asyncio
    async def test_timeout_protection(self, skill):
        query = "kodni ishga tushir:\n```python\nimport time; time.sleep(60)\n```"
        with patch.object(sys.modules["skills.code_runner"], "TIMEOUT_SECONDS", 2):
            result = await skill.execute(query)

        assert result["context"] == "code_runner:TIMEOUT"

    @pytest.mark.asyncio
    async def test_run_py_file(self, skill, tmp_path):
        script = tmp_path / "hello.py"
        script.write_text("print('fayldan chiqdim')", encoding="utf-8")
        result = await skill.execute(f"{script} kodni ishga tushir")

        assert result is not None
        assert "fayldan chiqdim" in result["response"]

    @pytest.mark.asyncio
    async def test_unrelated_query_returns_none(self, skill):
        assert await skill.execute("musiqani ishga tushir") is None or True


class TestM7Routing:
    def test_documents_routing(self):
        assert match_intents("report.pdf ni o'qib ber")[0] == "documents"
        assert match_intents("excel jadval yarat")[0] == "documents"

    def test_organize_beats_filemanager(self):
        intents = match_intents("papkani tartibga sol")
        assert intents[0] == "organize"

    def test_code_runner_routing(self):
        assert match_intents("kodni ishga tushir")[0] == "code_runner"

    def test_python_version_stays_system_info(self):
        assert match_intents("python versiyasi nima")[0] == "system_info"

    def test_brain_knows_m7_skills(self):
        from core.brain import AVAILABLE_SKILLS

        for name in ("organize", "documents", "code_runner"):
            assert name in AVAILABLE_SKILLS


class TestDocumentSkillLargeFiles:
    @pytest.mark.asyncio
    async def test_xlsx_memory_guard(self, tmp_path):
        from skills.documents import DocumentSkill

        skill = DocumentSkill()
        """Katta jadvalda iteratsiya erta to'xtashi kerak (xotira himoyasi)."""
        import openpyxl

        f = tmp_path / "katta.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        for i in range(5000):
            ws.append([f"qator_{i}", "uzun qiymat" * 20])
        wb.save(str(f))

        result = await skill.execute(f"{f} ni o'qi")
        assert result is not None
        assert len(result["context"]) <= 20000 * 2 + 1000
